from __future__ import annotations

import asyncio
import base64
import binascii
import logging
import os
import re
from io import BytesIO
from typing import Any, Dict
from urllib.parse import quote

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from dotenv import load_dotenv, find_dotenv
from fastapi import FastAPI, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
import yaml
from google import genai

from crewai import Agent, Task, Crew, LLM
from crewai_tools import SerperDevTool


logger = logging.getLogger(__name__)


# Simple config loader (YAML + env overrides)
DEFAULT_CONFIG_PATH = os.path.join(os.path.dirname(__file__), "..", "config", "config.yaml")


def load_config(path: str = DEFAULT_CONFIG_PATH) -> Dict[str, Any]:
    load_dotenv()  # Load .env from current or parent dirs
    load_dotenv(find_dotenv(), override=False)

    config = {}
    if os.path.exists(path):
        with open(path, "r", encoding="utf-8") as f:
            config = yaml.safe_load(f) or {}

    # Required env validation
    missing = []
    if not os.getenv("GOOGLE_API_KEY"):
        missing.append("GOOGLE_API_KEY")
    if not os.getenv("SERPER_API_KEY"):
        missing.append("SERPER_API_KEY")
    if missing:
        raise RuntimeError(f"Missing environment variables: {', '.join(missing)}")

    # Defaults with overrides
    app_cfg = config.get("app", {})
    llm_cfg = config.get("llm", {})
    crew_cfg = config.get("crew", {})

    return {
        "host": app_cfg.get("host", "127.0.0.1"),
        "port": app_cfg.get("port", 8002),
        "cors_origins": app_cfg.get("cors_origins", ["http://localhost:3000", "http://127.0.0.1:3000"]),
        "llm_model": llm_cfg.get("model", "gemini/gemini-3.5-flash"),
        "crew_verbose": crew_cfg.get("verbose", True),
    }


# FastAPI app
app = FastAPI()

# Load config early
settings = load_config()

# CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=settings["cors_origins"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


class TopicRequest(BaseModel):
    topic: str


class DocumentRequest(BaseModel):
    topic: str
    content: str
    image_data_url: str | None = None


def add_markdown_runs(paragraph: Any, text: str) -> None:
    text = re.sub(r"!\[([^]]*)\]\([^)]*\)", r"\1", text)
    text = re.sub(r"\[([^]]+)\]\([^)]+\)", r"\1", text)
    parts = re.split(r"(\*\*.+?\*\*|__.+?__|(?<!\*)\*[^*]+?\*(?!\*))", text)
    for part in parts:
        if not part:
            continue
        if (part.startswith("**") and part.endswith("**")) or (
            part.startswith("__") and part.endswith("__")
        ):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*"):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            paragraph.add_run(part.replace("`", ""))


def build_article_docx(request: DocumentRequest) -> BytesIO:
    document = Document()
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    heading_tokens = {
        "Heading 1": (16, "2E74B5", 16, 8),
        "Heading 2": (13, "2E74B5", 12, 6),
        "Heading 3": (12, "1F4D78", 8, 4),
    }
    for style_name, (size, color, before, after) in heading_tokens.items():
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    for style_name in ("List Bullet", "List Number"):
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing = 1.167

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(14)
    title_run = title.add_run(request.topic.strip() or "Generated Article")
    title_run.font.name = "Calibri"
    title_run.font.size = Pt(24)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor.from_string("0B2545")

    if request.image_data_url and request.image_data_url.startswith("data:image/"):
        try:
            _, encoded = request.image_data_url.split(",", 1)
            document.add_picture(BytesIO(base64.b64decode(encoded)), width=Inches(6.5))
            document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        except (ValueError, binascii.Error):
            logger.warning("Skipping invalid article image data")

    for raw_line in request.content.splitlines():
        line = raw_line.strip()
        if not line or line in {"---", "***", "___"}:
            continue

        heading_match = re.match(r"^(#{1,3})\s+(.+)$", line)
        bullet_match = re.match(r"^[-*+]\s+(.+)$", line)
        number_match = re.match(r"^\d+[.)]\s+(.+)$", line)

        if heading_match:
            paragraph = document.add_paragraph(style=f"Heading {len(heading_match.group(1))}")
            add_markdown_runs(paragraph, heading_match.group(2))
        elif bullet_match:
            paragraph = document.add_paragraph(style="List Bullet")
            add_markdown_runs(paragraph, bullet_match.group(1))
        elif number_match:
            paragraph = document.add_paragraph(style="List Number")
            add_markdown_runs(paragraph, number_match.group(1))
        else:
            paragraph = document.add_paragraph()
            add_markdown_runs(paragraph, line)

    output = BytesIO()
    document.save(output)
    output.seek(0)
    return output


def generate_topic_image(topic: str) -> str:
    client = genai.Client(api_key=os.getenv("GOOGLE_API_KEY"))
    interaction = client.interactions.create(
        model="gemini-3.1-flash-image",
        input=(
            "Create a polished, editorial-quality 16:9 hero image for a blog "
            f"article about: {topic}. Do not include logos, watermarks, or text."
        ),
        response_format={
            "type": "image",
            "mime_type": "image/jpeg",
            "aspect_ratio": "16:9",
            "image_size": "1K",
        },
    )
    image = interaction.output_image
    if image is None or not image.data:
        raise RuntimeError("Gemini returned no generated image")

    encoded = image.data
    if isinstance(encoded, bytes):
        encoded = base64.b64encode(encoded).decode("ascii")
    return f"{image.mime_type or 'image/jpeg'};base64,{encoded}"


# Crew builder (unchanged from notebook)
def build_crew() -> Crew:
    llm = LLM(
        api_key=os.getenv("GOOGLE_API_KEY"),
        model=settings["llm_model"],
    )

    os.environ["SERPER_API_KEY"] = os.getenv("SERPER_API_KEY", "")
    search_tool = SerperDevTool()

    planner = Agent(
        role="Content Planner",
        goal="Plan engaging and factually accurate content on {topic}",
        backstory=(
            "You're working on planning a blog article about the topic: {topic} in 'https://medium.com/'. "
            "You collect information that helps the audience learn something and make informed decisions. "
            "Prepare a detailed outline and the relevant topics and sub-topics for the blog post. "
            "Your work is the basis for the Content Writer."
        ),
        allow_delegation=False,
        verbose=settings["crew_verbose"],
        llm=llm,
        tools=[search_tool],
    )

    writer = Agent(
        role="Content Writer",
        goal="Write insightful and factually accurate opinion piece about the topic: {topic}",
        backstory=(
            "You're writing a new opinion piece about the topic: {topic} in 'https://medium.com/'. "
            "You base your writing on the Content Planner's outline, provide objective insights, and "
            "acknowledge when statements are opinions."
        ),
        allow_delegation=False,
        verbose=settings["crew_verbose"],
        llm=llm,
    )

    editor = Agent(
        role="Editor",
        goal=(
            "Edit a given blog post to align with the writing style of the organization 'https://medium.com/'."
        ),
        backstory=(
            "You review the blog post to ensure journalistic best practices, balanced viewpoints, and avoidance "
            "of major controversial topics when possible."
        ),
        allow_delegation=False,
        verbose=settings["crew_verbose"],
        llm=llm,
    )

    plan_task = Task(
        description=(
            "1. Prioritize the latest trends, key players, and noteworthy news on {topic}.\n"
            "2. Identify the target audience, considering their interests and pain points.\n"
            "3. Develop a detailed content outline including an introduction, key points, and a call to action.\n"
            "4. Include SEO keywords and relevant data or sources."
        ),
        expected_output=(
            "A comprehensive content plan with outline, audience analysis, SEO keywords, and resources."
        ),
        agent=planner,
        tools=[search_tool],
    )

    write_task = Task(
        description=(
            "1. Use the content plan to craft a compelling blog post on {topic}.\n"
            "2. Incorporate SEO keywords naturally.\n"
            "3. Sections/Subtitles are properly named in an engaging manner.\n"
            "4. Ensure the post has an engaging introduction, insightful body, and a summarizing conclusion.\n"
            "5. Proofread for grammatical errors and alignment with the brand's voice."
        ),
        expected_output=(
            "A well-written blog post in markdown format, ready for publication, with 2-3 paragraphs per section."
        ),
        agent=writer,
    )

    edit_task = Task(
        description=(
            "Proofread the given blog post for grammatical errors and alignment with the brand's voice."
        ),
        expected_output=(
            "A well-written blog post in markdown format (no leading word 'markdown'), ready for publication, "
            "with 2-3 paragraphs per section."
        ),
        agent=editor,
    )

    return Crew(
        agents=[planner, writer, editor],
        tasks=[plan_task, write_task, edit_task],
        verbose=settings["crew_verbose"],
    )


# Store crew in app state at startup
@app.on_event("startup")
def startup() -> None:
    app.state.crew = build_crew()


@app.post("/generate-blog/")
async def generate_blog(request: TopicRequest) -> Dict[str, Any]:
    if not request.topic or not request.topic.strip():
        raise HTTPException(status_code=400, detail="'topic' must be provided")

    crew = app.state.crew
    try:
        # CrewAI and some of its tools use synchronous internals. Running the
        # workflow in a worker thread keeps them outside FastAPI's event loop.
        result = await asyncio.to_thread(
            crew.kickoff,
            inputs={"topic": request.topic.strip()},
        )
        blog_text = getattr(result, "raw", None) or str(result)
        return {"topic": request.topic, "blog": {"raw": blog_text}}
    except Exception as exc:
        logger.exception("Blog generation failed for topic %r", request.topic)
        raise HTTPException(status_code=500, detail=str(exc)) from exc


@app.post("/generate-image/")
async def generate_image(request: TopicRequest) -> Dict[str, str]:
    topic = request.topic.strip()
    if not topic:
        raise HTTPException(status_code=400, detail="'topic' must be provided")

    try:
        image_data = await asyncio.to_thread(generate_topic_image, topic)
        return {"imageUrl": f"data:{image_data}"}
    except Exception as exc:
        logger.exception("Image generation failed for topic %r", topic)
        raise HTTPException(status_code=500, detail=str(exc)) from exc


@app.post("/download-word/")
async def download_word(request: DocumentRequest) -> StreamingResponse:
    if not request.content.strip():
        raise HTTPException(status_code=400, detail="Article content must be provided")

    try:
        document = await asyncio.to_thread(build_article_docx, request)
        safe_stem = re.sub(r"[^A-Za-z0-9_-]+", "-", request.topic).strip("-")
        filename = f"{safe_stem or 'generated-article'}.docx"
        encoded_filename = quote(filename)
        return StreamingResponse(
            document,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={
                "Content-Disposition": (
                    f'attachment; filename="{filename}"; filename*=UTF-8\'\'{encoded_filename}'
                )
            },
        )
    except Exception as exc:
        logger.exception("Word document creation failed for topic %r", request.topic)
        raise HTTPException(status_code=500, detail=str(exc)) from exc


